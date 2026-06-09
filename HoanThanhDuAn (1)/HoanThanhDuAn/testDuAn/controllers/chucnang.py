# ===== IMPORT THƯ VIỆN =====
import tkinter as tk
from tkinter import messagebox, filedialog
import threading


# ===== HÀM XÓA Ô NHẬP =====
def xoa_o_nhap(
    entry_ten,
    entry_loai,
    entry_nguyenlieu,
    entry_dinhluong,
    entry_thoigian,
    entry_danhgia
):

    # Xóa tên món
    entry_ten.delete(0, tk.END)

    # Xóa loại món
    entry_loai.delete(0, tk.END)

    # Xóa nguyên liệu
    entry_nguyenlieu.delete(0, tk.END)

    # Xóa định lượng
    entry_dinhluong.delete(0, tk.END)

    # Xóa thời gian
    entry_thoigian.delete(0, tk.END)

    # Xóa đánh giá
    entry_danhgia.delete(0, tk.END)


# ===== HÀM THÊM MÓN =====
def them_mon(
    tree,
    entry_ten,
    entry_loai,
    entry_nguyenlieu,
    entry_dinhluong,
    entry_thoigian,
    entry_danhgia
):

    # Lấy dữ liệu
    ten = entry_ten.get().strip()
    loai = entry_loai.get().strip()
    nguyenlieu = entry_nguyenlieu.get().strip()
    dinhluong = entry_dinhluong.get().strip()
    thoigian = entry_thoigian.get().strip()
    danhgia = entry_danhgia.get().strip()

    # Kiểm tra rỗng
    if (
        ten == "" or
        loai == "" or
        nguyenlieu == "" or
        dinhluong == "" or
        thoigian == "" or
        danhgia == ""
    ):

        messagebox.showwarning(
            "Thông báo",
            "Vui lòng nhập đầy đủ thông tin"
        )

        return

    # Kiểm tra thời gian
    if not thoigian.isdigit():

        messagebox.showwarning(
            "Thông báo",
            "Thời gian phải là số"
        )

        return

    # Thêm dữ liệu vào bảng
    tree.insert(
        "",
        tk.END,
        values=(
            "☐",
            ten,
            loai,
            nguyenlieu,
            dinhluong,
            thoigian,
            danhgia
        )
    )

    # Xóa ô nhập
    xoa_o_nhap(
        entry_ten,
        entry_loai,
        entry_nguyenlieu,
        entry_dinhluong,
        entry_thoigian,
        entry_danhgia
    )

    # Thông báo
    messagebox.showinfo(
        "Thông báo",
        "Thêm món thành công"
    )


# ===== HÀM XÓA MÓN =====
def xoa_mon(
    tree,
    entry_ten,
    entry_loai,
    entry_nguyenlieu,
    entry_dinhluong,
    entry_thoigian,
    entry_danhgia
):

    # Lấy tất cả dòng có tích chọn
    checked_items = []
    for item in tree.get_children():
        vals = tree.item(item, "values")
        if vals and vals[0] == "☑":
            checked_items.append(item)

    # Nếu không có dòng nào được tích chọn, sử dụng dòng đang được chọn highlight
    if not checked_items:
        selected = tree.selection()
        if selected:
            checked_items = list(selected)

    # Kiểm tra xem có gì để xóa không
    if not checked_items:

        messagebox.showwarning(
            "Thông báo",
            "Vui lòng tích chọn hoặc chọn dòng cần xóa"
        )

        return

    # Xác nhận xóa nhiều món nếu số lượng lớn hơn 1
    if len(checked_items) > 1:
        if not messagebox.askyesno("Xác nhận", f"Bạn có chắc chắn muốn xóa {len(checked_items)} món đã chọn?"):
            return
    else:
        # Xác nhận xóa 1 món
        if not messagebox.askyesno("Xác nhận", "Bạn có chắc chắn muốn xóa món đã chọn?"):
            return

    # Xóa dòng
    for item in checked_items:

        tree.delete(item)

    # Xóa ô nhập
    xoa_o_nhap(
        entry_ten,
        entry_loai,
        entry_nguyenlieu,
        entry_dinhluong,
        entry_thoigian,
        entry_danhgia
    )

    # Bỏ chọn dòng
    tree.selection_remove(tree.selection())

    # Cập nhật bảng
    tree.update_idletasks()

    # Thông báo
    messagebox.showinfo(
        "Thông báo",
        "Xóa món thành công"
    )


# ===== HÀM SỬA MÓN =====
def sua_mon(
    tree,
    entry_ten,
    entry_loai,
    entry_nguyenlieu,
    entry_dinhluong,
    entry_thoigian,
    entry_danhgia
):

    # Lấy dòng đang chọn
    selected = tree.selection()

    # Kiểm tra chọn dòng
    if not selected:

        messagebox.showwarning(
            "Thông báo",
            "Chọn món cần sửa"
        )

        return

    # Lấy dữ liệu
    ten = entry_ten.get().strip()
    loai = entry_loai.get().strip()
    nguyenlieu = entry_nguyenlieu.get().strip()
    dinhluong = entry_dinhluong.get().strip()
    thoigian = entry_thoigian.get().strip()
    danhgia = entry_danhgia.get().strip()

    # Kiểm tra rỗng
    if (
        ten == "" or
        loai == "" or
        nguyenlieu == "" or
        dinhluong == "" or
        thoigian == "" or
        danhgia == ""
    ):

        messagebox.showwarning(
            "Thông báo",
            "Vui lòng nhập đầy đủ thông tin"
        )

        return

    # Kiểm tra thời gian
    if not thoigian.isdigit():

        messagebox.showwarning(
            "Thông báo",
            "Thời gian phải là số"
        )

        return

    # Cập nhật dữ liệu
    current_values = tree.item(selected[0], "values")
    checkbox_val = current_values[0] if current_values else "☐"
    tree.item(
        selected[0],
        values=(
            checkbox_val,
            ten,
            loai,
            nguyenlieu,
            dinhluong,
            thoigian,
            danhgia
        )
    )

    # Xóa ô nhập
    xoa_o_nhap(
        entry_ten,
        entry_loai,
        entry_nguyenlieu,
        entry_dinhluong,
        entry_thoigian,
        entry_danhgia
    )

    # Bỏ chọn dòng
    tree.selection_remove(tree.selection())

    # Cập nhật bảng
    tree.update_idletasks()

    # Thông báo
    messagebox.showinfo(
        "Thông báo",
        "Sửa món thành công"
    )


# ===== HÀM CHỌN MÓN =====
def chon_mon(
    event,
    tree,
    entry_ten,
    entry_loai,
    entry_nguyenlieu,
    entry_dinhluong,
    entry_thoigian,
    entry_danhgia
):

    # Lấy dòng đang chọn
    selected = tree.selection()

    # Nếu có dòng được chọn
    if selected:

        # Lấy dữ liệu dòng
        values = tree.item(selected[0], "values")

        # Hiện tên món
        entry_ten.delete(0, tk.END)
        entry_ten.insert(0, values[1])

        # Hiện loại món
        entry_loai.delete(0, tk.END)
        entry_loai.insert(0, values[2])

        # Hiện nguyên liệu
        entry_nguyenlieu.delete(0, tk.END)
        entry_nguyenlieu.insert(0, values[3])

        # Hiện định lượng
        entry_dinhluong.delete(0, tk.END)
        entry_dinhluong.insert(0, values[4])

        # Hiện thời gian
        entry_thoigian.delete(0, tk.END)
        entry_thoigian.insert(0, values[5])

        # Hiện đánh giá
        entry_danhgia.delete(0, tk.END)
        entry_danhgia.insert(0, values[6])


# ===== HÀM XUẤT WORD =====
def xuat_word(tree):
    # Lấy tất cả dòng từ treeview
    children = tree.get_children()
    if not children:
        messagebox.showwarning(
            "Thông báo",
            "Danh sách trống, không có gì để xuất!"
        )
        return

    # Lấy tất cả dữ liệu từ treeview ngay trên luồng chính để đảm bảo thread-safety
    # Bỏ qua cột "Chọn" đầu tiên bằng slice [1:]
    data_list = [tree.item(child, "values")[1:] for child in children]

    # Chọn vị trí và tên file để lưu
    filepath = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
        title="Chọn nơi lưu file Word"
    )

    # Nếu người dùng hủy
    if not filepath:
        return

    # Lấy cửa sổ cha và đặt con trỏ chuột sang trạng thái chờ "watch"
    root = tree.winfo_toplevel()
    root.config(cursor="watch")

    # Các hàm callback để gọi lại trên luồng chính khi luồng phụ chạy xong
    def export_success():
        root.config(cursor="")
        messagebox.showinfo(
            "Thông báo",
            "Xuất file Word thành công!"
        )

    def export_error(err_msg):
        root.config(cursor="")
        messagebox.showerror(
            "Lỗi",
            f"Không thể xuất file Word: {err_msg}"
        )

    # Định nghĩa hàm xử lý ngầm (Worker function)
    def worker():
        try:
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
            except ImportError:
                import subprocess
                import sys
                subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn

            # Tạo document mới
            doc = Document()

            # Cấu hình font mặc định
            style = doc.styles['Normal']
            style.font.name = 'Segoe UI'
            style.font.size = Pt(11)
            style.font.color.rgb = RGBColor(51, 65, 85) # Slate #334155

            # Tiêu đề tài liệu
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run("DANH SÁCH CÔNG THỨC NẤU ĂN")
            title_run.font.name = 'Segoe UI'
            title_run.font.size = Pt(22)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(30, 41, 59) # Slate đậm #1E293B
            title_p.paragraph_format.space_after = Pt(4)

            # Mô tả / Phụ đề
            sub_p = doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_p.add_run("Được xuất từ phần mềm Quản Lý Công Thức Nấu Ăn")
            sub_run.font.name = 'Segoe UI'
            sub_run.font.size = Pt(11)
            sub_run.font.italic = True
            sub_run.font.color.rgb = RGBColor(100, 116, 139) # Gray #64748B
            sub_p.paragraph_format.space_after = Pt(24)

            # Tiêu đề các cột
            headers = ["Tên món", "Loại món", "Nguyên liệu", "Định lượng", "Thời gian (phút)", "Đánh giá"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'

            # Hàm helper đổ màu nền cho ô trong table
            def set_cell_background(cell, hex_color):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), hex_color)
                tcPr.append(shd)

            # Định dạng Header của bảng
            hdr_cells = table.rows[0].cells
            for i, header_text in enumerate(headers):
                hdr_cells[i].text = header_text
                set_cell_background(hdr_cells[i], "1E293B") # Màu Slate #1E293B làm chủ đạo
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Segoe UI'
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255) # Chữ trắng nổi bật

            # Điền dữ liệu từ danh sách
            for r_idx, row_values in enumerate(data_list):
                row_cells = table.add_row().cells

                # Sử dụng hiệu ứng Zebra striping (đổ màu nền xen kẽ)
                bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"

                for c_idx, val in enumerate(row_values):
                    row_cells[c_idx].text = str(val)
                    set_cell_background(row_cells[c_idx], bg_color)

                    for paragraph in row_cells[c_idx].paragraphs:
                        # Canh lề cột thời gian và đánh giá ra giữa, các cột còn lại canh trái
                        if c_idx in [4, 5]:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        for run in paragraph.runs:
                            run.font.name = 'Segoe UI'
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(51, 65, 85)

            # Phần tổng kết ở cuối tài liệu
            doc.add_paragraph().paragraph_format.space_before = Pt(20)
            summary_p = doc.add_paragraph()
            summary_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            summary_run = summary_p.add_run(f"Tổng số lượng: {len(data_list)} công thức nấu ăn.")
            summary_run.font.name = 'Segoe UI'
            summary_run.font.bold = True
            summary_run.font.size = Pt(11)
            summary_run.font.color.rgb = RGBColor(37, 99, 235) # Xanh Royal đậm chất công nghệ #2563EB

            # Lưu file
            doc.save(filepath)
            
            # Gọi callback thành công một cách an toàn trên luồng chính
            root.after(0, export_success)

        except Exception as e:
            # Gọi callback báo lỗi một cách an toàn trên luồng chính
            root.after(0, lambda err=e: export_error(str(err)))

    # Tạo và chạy luồng phụ (background thread)
    t = threading.Thread(target=worker, daemon=True)
    t.start()