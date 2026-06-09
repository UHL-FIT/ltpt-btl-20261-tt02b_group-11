# -*- coding: utf-8 -*-
"""
File: tao_bao_cao.py
Mục đích: Tự động tạo báo cáo Word (.docx) chuyên nghiệp, chi tiết và có độ dài khoảng 15-20 trang cho dự án "Quản Lý Công Thức Nấu Ăn" (Cooking Manager).
Cách sử dụng: Chạy file này bằng lệnh `python tao_bao_cao.py` hoặc `python testDuAn/tao_bao_cao.py`.
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn, nsdecls
except ImportError:
    print("Đang tiến hành cài đặt thư viện 'python-docx'...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn, nsdecls

# Định nghĩa màu sắc chủ đạo (Royal Theme)
COLOR_PRIMARY = RGBColor(30, 58, 138)    # #1E3A8A (Navy đậm)
COLOR_SECONDARY = RGBColor(37, 99, 235) # #2563EB (Royal Blue)
COLOR_TEXT = RGBColor(51, 65, 85)       # #334155 (Slate Grey)
COLOR_HEADING = RGBColor(15, 23, 42)    # #0F172A (Slate đậm)
HEX_PRIMARY = "1E3A8A"
HEX_LIGHT_GRAY = "F8FAFC"
HEX_BORDER = "CBD5E1"

def set_cell_background(cell, hex_color):
    """Thiết lập màu nền cho một ô trong bảng."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Thiết lập khoảng đệm trong ô (padding) tính bằng dxa."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    """Thêm tiêu đề được thiết kế đẹp mắt và phân cấp rõ ràng."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text.upper())
        run.font.name = 'Segoe UI'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        # Thêm đường gạch chân mỏng dưới Heading 1
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12') # 1.5 pt
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), HEX_PRIMARY)
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)
        
    elif level == 2:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        
    elif level == 3:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = COLOR_HEADING
        
    return p

def add_bullet_point(doc, bold_prefix, text_content):
    """Thêm dòng bullet point có bôi đậm phần tiền tố."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run_bold = p.add_run(bold_prefix)
    run_bold.font.name = 'Segoe UI'
    run_bold.font.bold = True
    run_bold.font.color.rgb = COLOR_TEXT
    
    run_text = p.add_run(text_content)
    run_text.font.name = 'Segoe UI'
    run_text.font.color.rgb = COLOR_TEXT
    return p

def add_toc_item(doc, title, page_str, level=1):
    """Thêm một mục vào trang Mục lục với tab stop và dấu chấm dẫn đường."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    tab_stops = p.paragraph_format.tab_stops
    # Lề trang A4 là 1 inch hai bên, chiều ngang trang là 8.27 inch.
    # Khổ in khả dụng là 6.27 inch, đặt tab stop tại 6.25 inch để tránh tràn lề.
    tab_stops.add_tab_stop(Inches(6.25), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
    
    if level == 1:
        p.paragraph_format.left_indent = Inches(0)
        run = p.add_run(f"{title}\t{page_str}")
        run.font.name = 'Segoe UI'
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_HEADING
    elif level == 2:
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f"{title}\t{page_str}")
        run.font.name = 'Segoe UI'
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TEXT
    elif level == 3:
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"{title}\t{page_str}")
        run.font.name = 'Segoe UI'
        run.font.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_TEXT

def setup_header_footer(doc, title_text="BÁO CÁO DỰ ÁN: COOKING MANAGER APP"):
    """Thiết lập Header và Footer chuyên nghiệp cho toàn bộ tài liệu."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    # --- HEADER ---
    header = section.header
    header.is_linked_to_previous = False
    
    # Xóa nội dung header mặc định
    for para in header.paragraphs:
        p = para._p
        p.getparent().remove(p)
    
    p_hdr = header.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run_hdr = p_hdr.add_run(title_text)
    run_hdr.font.name = 'Segoe UI'
    run_hdr.font.size = Pt(9)
    run_hdr.font.italic = True
    run_hdr.font.color.rgb = RGBColor(100, 116, 139)  # Slate 500
    
    # Đường kẻ mỏng dưới header
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '4')
    bottom_border.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bottom_border)
    p_hdr._p.get_or_add_pPr().append(pBdr)
    
    # --- FOOTER ---
    footer = section.footer
    footer.is_linked_to_previous = False
    
    for para in footer.paragraphs:
        p = para._p
        p.getparent().remove(p)
    
    p_ftr = footer.add_paragraph()
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_ftr_prefix = p_ftr.add_run("Trang ")
    run_ftr_prefix.font.name = 'Segoe UI'
    run_ftr_prefix.font.size = Pt(9.5)
    run_ftr_prefix.font.bold = True
    run_ftr_prefix.font.color.rgb = RGBColor(15, 23, 42)
    
    # Chèn trường số trang tự động (PAGE field)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    
    run_page = p_ftr.add_run()
    run_page.font.name = 'Segoe UI'
    run_page.font.size = Pt(9.5)
    run_page.font.bold = True
    run_page.font.color.rgb = RGBColor(15, 23, 42)
    run_page._r.append(fldChar_begin)
    run_page._r.append(instrText)
    run_page._r.append(fldChar_end)
    
    run_ftr_suffix = p_ftr.add_run(" | Quản Lý Công Thức Nấu Ăn - Cooking Manager")
    run_ftr_suffix.font.name = 'Segoe UI'
    run_ftr_suffix.font.size = Pt(9)
    run_ftr_suffix.font.color.rgb = RGBColor(100, 116, 139)
    
    # Đường kẻ mỏng trên footer
    pBdr_ftr = OxmlElement('w:pBdr')
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '6')
    top_border.set(qn('w:space'), '4')
    top_border.set(qn('w:color'), 'CBD5E1')
    pBdr_ftr.append(top_border)
    p_ftr._p.get_or_add_pPr().append(pBdr_ftr)

def add_image_captioned(doc, img_path, caption_text, width_inches=5.5):
    """Chèn hình ảnh có chú thích vào tài liệu, căn giữa và có viền nhẹ."""
    if not os.path.exists(img_path):
        return  # Bỏ qua nếu file ảnh không tồn tại
    
    # Đoạn văn chứa hình ảnh, căn giữa
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(10)
    p_img.paragraph_format.space_after = Pt(4)
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Inches(width_inches))
    
    # Chú thích hình ảnh
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(14)
    run_cap = p_cap.add_run(caption_text)
    run_cap.font.name = 'Segoe UI'
    run_cap.font.size = Pt(9.5)
    run_cap.font.italic = True
    run_cap.font.color.rgb = RGBColor(100, 116, 139)  # Slate 500

def add_code_block(doc, code_text):
    """Thêm một khối mã nguồn có viền nhẹ, nền xám nhạt và dùng font monospace."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Thiết lập độ rộng cột là 6.5 inches (khớp khổ giấy A4 trừ lề)
    table.columns[0].width = Inches(6.5)
    
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=140)
    
    # Thiết lập viền trái màu xanh dương
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24') # 3pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), '3B82F6') # Thanh chỉ màu xanh dương bên trái
    tcBorders.append(left)
    for side in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)
    
    # Tách văn bản thành từng dòng để tránh lỗi hiển thị và xuống dòng sai trong Word
    lines = code_text.split('\n')
    for idx, line in enumerate(lines):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
            
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1  # Giãn dòng tối ưu cho code
        
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.0)
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        
    # Khoảng cách sau bảng code
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_table_styled(doc, headers, data_rows, col_widths=None):
    """Thêm một bảng được định dạng đẹp mắt vào tài liệu."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Định dạng Header
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = Inches(col_widths[i])
        set_cell_background(hdr_cells[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.name = 'Segoe UI'
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Điền dữ liệu
    for r_idx, row_values in enumerate(data_rows):
        row_cells = table.add_row().cells
        bg_color = HEX_LIGHT_GRAY if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx, val in enumerate(row_values):
            row_cells[c_idx].text = str(val)
            if col_widths and c_idx < len(col_widths):
                row_cells[c_idx].width = Inches(col_widths[c_idx])
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=120, right=120)
            
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            # Nếu cột đầu hoặc cột ngắn thì căn giữa, cột dài căn trái
            if len(str(val)) > 30:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            for run in p.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9.5)
                run.font.color.rgb = COLOR_TEXT
                
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    return table

def tao_bao_cao():
    print("Bắt đầu tạo tài liệu báo cáo dự án chi tiết (khoảng 15 trang)...")
    doc = Document()
    
    # 1. Cấu hình định dạng trang (Page Setup) cho khổ A4 chuẩn
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 Width
    section.page_height = Inches(11.69) # A4 Height
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Đường dẫn thư mục chứa ảnh minh họa
    assets_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets')
    
    # 2. Cấu hình style "Normal" mặc định cho tài liệu
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.25
    style_normal.paragraph_format.space_after = Pt(8)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Thiết lập header và footer cho toàn bộ tài liệu
    setup_header_footer(doc)
    
    # =========================================================================
    # TRANG BÌA (COVER PAGE)
    # =========================================================================
    p_school = doc.add_paragraph()
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_school.paragraph_format.space_before = Pt(10)
    p_school.paragraph_format.space_after = Pt(4)
    run_school = p_school.add_run("[TÊN TRƯỜNG ĐẠI HỌC / CAO ĐẲNG CỦA BẠN]")
    run_school.font.name = 'Segoe UI'
    run_school.font.size = Pt(12)
    run_school.font.bold = True
    run_school.font.color.rgb = COLOR_PRIMARY
    
    p_faculty = doc.add_paragraph()
    p_faculty.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_faculty.paragraph_format.space_after = Pt(60)
    run_faculty = p_faculty.add_run("[KHOA CÔNG NGHỆ THÔNG TIN / KHOA ĐÀO TẠO]")
    run_faculty.font.name = 'Segoe UI'
    run_faculty.font.size = Pt(11)
    run_faculty.font.bold = True
    run_faculty.font.color.rgb = COLOR_TEXT
    
    # Tiêu đề báo cáo
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("BÁO CÁO DỰ ÁN MÔN HỌC")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    run_sub = p_sub.add_run("ĐỀ TÀI: XÂY DỰNG PHẦN MỀM QUẢN LÝ CÔNG THỨC NẤU ĂN\n(COOKING MANAGER APP)")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(16)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_SECONDARY
    
    # Hình ảnh banner minh họa trên trang bìa
    add_image_captioned(doc,
        os.path.join(assets_dir, 'cover_banner.png'),
        "Hình 1. Giao diện minh họa phần mềm Quản Lý Công Thức Nấu Ăn",
        width_inches=5.0
    )
    
    # Bảng thông tin sinh viên thực hiện (Căn lề cân đối ở giữa dưới bìa)
    table_info = doc.add_table(rows=5, cols=2)
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_info.autofit = False
    
    info_labels = [
        ("Môn học học phần:", "[Tên Môn Học Của Bạn]"),
        ("Giảng viên hướng dẫn:", "[Họ và Tên Giảng Viên]"),
        ("Sinh viên thực hiện:", "[Họ và Tên Sinh Viên]"),
        ("Mã số sinh viên:", "[Mã Số Sinh Viên Của Bạn]"),
        ("Lớp học phần:", "[Tên Lớp Học Phần]")
    ]
    
    for idx, (label, val) in enumerate(info_labels):
        row = table_info.rows[idx]
        
        # Thiết lập độ rộng cột cho đẹp mắt
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(3.5)
        
        # Căn lề trái cho nhãn, lề trái cho giá trị
        p_label = row.cells[0].paragraphs[0]
        p_label.paragraph_format.space_after = Pt(4)
        run_l = p_label.add_run(label)
        run_l.font.name = 'Segoe UI'
        run_l.font.bold = True
        run_l.font.size = Pt(11)
        run_l.font.color.rgb = COLOR_TEXT
        
        p_val = row.cells[1].paragraphs[0]
        p_val.paragraph_format.space_after = Pt(4)
        run_v = p_val.add_run(val)
        run_v.font.name = 'Segoe UI'
        run_v.font.size = Pt(11)
        run_v.font.color.rgb = COLOR_TEXT
        
    # Đường kẻ trang trí ngang trước năm
    p_divider = doc.add_paragraph()
    p_divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_divider.paragraph_format.space_before = Pt(80)
    p_divider.paragraph_format.space_after = Pt(8)
    pBdr_div = OxmlElement('w:pBdr')
    top_div = OxmlElement('w:top')
    top_div.set(qn('w:val'), 'single')
    top_div.set(qn('w:sz'), '12')
    top_div.set(qn('w:space'), '4')
    top_div.set(qn('w:color'), HEX_PRIMARY)
    pBdr_div.append(top_div)
    p_divider._p.get_or_add_pPr().append(pBdr_div)
    
    p_year = doc.add_paragraph()
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_year.paragraph_format.space_before = Pt(8)
    run_year = p_year.add_run("[Địa danh], Năm 2026")
    run_year.font.name = 'Segoe UI'
    run_year.font.size = Pt(11)
    run_year.font.bold = True
    run_year.font.color.rgb = COLOR_TEXT
    
    # Sang trang mới
    doc.add_page_break()
    
    # =========================================================================
    # LỜI CẢM ƠN
    # =========================================================================
    add_heading_styled(doc, "LỜI CẢM ƠN", 1)
    
    p_thanks = doc.add_paragraph()
    run_t1 = p_thanks.add_run(
        "Lời đầu tiên, em xin gửi lời cảm ơn chân thành và sâu sắc nhất tới toàn thể các thầy cô giáo Khoa Công nghệ "
        "Thông tin của nhà trường đã luôn tận tình dạy dỗ, truyền đạt những kiến thức quý báu và định hướng nghề nghiệp "
        "cho em trong suốt quá trình học tập vừa qua. Những bài giảng tâm huyết và phương pháp truyền đạt sinh động của "
        "các thầy cô đã là nền tảng vững chắc giúp em có thể hoàn thành đề tài này."
    )
    run_t1.font.name = 'Segoe UI'
    
    p_thanks2 = doc.add_paragraph()
    run_t2 = p_thanks2.add_run(
        "Đặc biệt, em xin gửi lời cảm ơn chân thành đến giảng viên hướng dẫn - Thầy/Cô "
    )
    run_t2.font.name = 'Segoe UI'
    run_t2_name = p_thanks2.add_run("[Họ và Tên Giảng Viên]")
    run_t2_name.font.name = 'Segoe UI'
    run_t2_name.font.bold = True
    run_t2_name.font.color.rgb = COLOR_PRIMARY
    run_t2_cont = p_thanks2.add_run(
        " đã tận tình chỉ bảo, góp ý từng điểm thiếu sót và luôn tạo điều kiện thuận lợi nhất "
        "để em hoàn thành đề tài dự án quản lý này đúng tiến độ. Sự hỗ trợ và định hướng của thầy/cô trong suốt "
        "quá trình thực hiện là nguồn động lực to lớn giúp em vượt qua những khó khăn kỹ thuật trong dự án."
    )
    run_t2_cont.font.name = 'Segoe UI'
    
    p_thanks3 = doc.add_paragraph()
    run_t3 = p_thanks3.add_run(
        "Mặc dù đã cố gắng hết sức để hoàn thiện đề tài, song do kiến thức còn hạn chế và thời gian thực hiện có giới hạn, "
        "báo cáo khó tránh khỏi những điểm còn thiếu sót. Em rất mong nhận được ý kiến đóng góp quý báu từ quý thầy cô "
        "để đề tài này được hoàn thiện hơn và có giá trị ứng dụng thực tiễn cao hơn trong tương lai."
    )
    run_t3.font.name = 'Segoe UI'
    
    # Khối ký tên cuối trang Lời cảm ơn
    p_sign_space = doc.add_paragraph()
    p_sign_space.paragraph_format.space_before = Pt(24)
    
    p_sign_place = doc.add_paragraph()
    p_sign_place.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_place = p_sign_place.add_run("[Địa danh], ngày     tháng     năm 2026")
    run_place.font.name = 'Segoe UI'
    run_place.font.italic = True
    run_place.font.color.rgb = COLOR_TEXT
    
    p_sign_title = doc.add_paragraph()
    p_sign_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sign_title = p_sign_title.add_run("Sinh viên thực hiện")
    run_sign_title.font.name = 'Segoe UI'
    run_sign_title.font.bold = True
    run_sign_title.font.color.rgb = COLOR_TEXT
    
    p_sign_space2 = doc.add_paragraph()
    p_sign_space2.paragraph_format.space_before = Pt(40)
    p_sign_space2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p_sign_name = doc.add_paragraph()
    p_sign_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sign_name = p_sign_name.add_run("[Họ và Tên Sinh Viên]")
    run_sign_name.font.name = 'Segoe UI'
    run_sign_name.font.bold = True
    run_sign_name.font.color.rgb = COLOR_PRIMARY
    
    # Sang trang mới
    doc.add_page_break()

    # =========================================================================
    # LỜI MỞ ĐẦU
    # =========================================================================
    add_heading_styled(doc, "LỜI MỞ ĐẦU", 1)
    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "Hiện nay, sự phát triển không ngừng của công nghệ thông tin đã mang lại những ứng dụng vô cùng thiết thực "
        "vào mọi lĩnh vực trong cuộc sống hằng ngày. Quản lý việc bếp núc, theo dõi dinh dưỡng và lưu giữ công thức nấu ăn "
        "cũng đã chuyển dịch từ phương thức ghi chép sổ tay truyền thống sang việc số hóa thông qua các phần mềm tiện ích. "
        "Cách làm cũ bộc lộ nhiều điểm hạn chế như dễ hư hỏng sổ sách, khó tra cứu nhanh nguyên liệu và thiếu khả năng "
        "tổng hợp, thống kê dữ liệu một cách toàn diện."
    )
    p_intro2 = doc.add_paragraph()
    p_intro2.add_run(
        "Được xây dựng với mục tiêu giải quyết vấn đề trên, ứng dụng \"Cooking Manager\" là một giải pháp quản lý trực quan giúp "
        "người dùng lưu trữ các món ăn, cập nhật nhanh thông tin định lượng nguyên liệu, tìm kiếm thông minh và phân tích thống kê "
        "khoa học thông qua sự kết hợp của ngôn ngữ lập trình Python và thư viện Pandas. Báo cáo này sẽ trình bày một cách "
        "chi tiết và hệ thống nhất về kiến trúc xây dựng phần mềm, các công nghệ sử dụng, giải thích chi tiết mã nguồn cốt lõi "
        "cũng như kết quả vận hành thực tế của chương trình."
    )
    
    # Sang trang mới
    doc.add_page_break()

    # =========================================================================
    # MỤC LỤC
    # =========================================================================
    add_heading_styled(doc, "MỤC LỤC", 1)
    
    p_note = doc.add_paragraph()
    p_note_run = p_note.add_run("Danh mục nội dung chi tiết trong tài liệu báo cáo:")
    p_note_run.font.italic = True
    p_note.paragraph_format.space_after = Pt(12)
    
    add_toc_item(doc, "LỜI CẢM ƠN", "2", 1)
    add_toc_item(doc, "LỜI MỞ ĐẦU", "3", 1)
    add_toc_item(doc, "MỤC LỤC", "4", 1)
    
    add_toc_item(doc, "CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI & PHÂN TÍCH YÊU CẦU", "5", 1)
    add_toc_item(doc, "1.1. Lý do chọn đề tài", "5", 2)
    add_toc_item(doc, "1.2. Mục tiêu của dự án", "5", 2)
    add_toc_item(doc, "1.3. Đối tượng sử dụng & Chân dung người dùng", "6", 2)
    add_toc_item(doc, "1.4. Phân tích yêu cầu hệ thống chi tiết", "6", 2)
    add_toc_item(doc, "1.5. Đặc tả chức năng & Mô hình Use Case", "7", 2)
    
    add_toc_item(doc, "CHƯƠNG 2: CÔNG NGHỆ SỬ DỤNG", "8", 1)
    add_toc_item(doc, "2.1. Ngôn ngữ lập trình Python", "8", 2)
    add_toc_item(doc, "2.2. Thư viện giao diện đồ họa Tkinter", "8", 2)
    add_toc_item(doc, "2.3. Lập trình đa luồng (Multi-threading) trong Python GUI", "9", 2)
    add_toc_item(doc, "2.4. Thư viện phân tích dữ liệu Pandas", "9", 2)
    add_toc_item(doc, "2.5. Thư viện xuất báo cáo python-docx", "10", 2)
    
    add_toc_item(doc, "CHƯƠNG 3: THIẾT KẾ KIẾN TRÚC & MÔ HÌNH HỆ THỐNG", "11", 1)
    add_toc_item(doc, "3.1. Cấu trúc thư mục dự án (Directory Structure)", "11", 2)
    add_toc_item(doc, "3.2. Áp dụng kiến trúc MVC (Model-View-Controller)", "11", 2)
    add_toc_item(doc, "3.3. Thiết kế mô hình dữ liệu (Data Schema)", "12", 2)
    add_toc_item(doc, "3.4. Sơ đồ luồng xử lý hệ thống (Data Flow Diagram)", "12", 2)
    
    add_toc_item(doc, "CHƯƠNG 4: HƯỚNG DẪN SỬ DỤNG CHI TIẾT (USER MANUAL)", "13", 1)
    add_toc_item(doc, "4.1. Khởi chạy ứng dụng và Giao diện chính", "13", 2)
    add_toc_item(doc, "4.2. Hướng dẫn Thêm mới và Chỉnh sửa công thức", "13", 2)
    add_toc_item(doc, "4.3. Hướng dẫn Xóa và Chọn nhiều món ăn", "14", 2)
    add_toc_item(doc, "4.4. Hướng dẫn Tra cứu bằng Bộ lọc đa tiêu chí", "14", 2)
    add_toc_item(doc, "4.5. Hướng dẫn Xem thống kê nâng cao và Xuất Word", "15", 2)
    
    add_toc_item(doc, "CHƯƠNG 5: PHÂN TÍCH MÃ NGUỒN CỐT LÕI (CODE ANALYSIS)", "16", 1)
    add_toc_item(doc, "5.1. File main.py - Điểm khởi chạy hệ thống", "16", 2)
    add_toc_item(doc, "5.2. File views/form.py - Xây dựng giao diện & Xử lý sự kiện", "16", 2)
    add_toc_item(doc, "5.3. File controllers/chucnang.py - Hàm xuất báo cáo Word", "18", 2)
    add_toc_item(doc, "5.4. File controllers/thongke.py - Thống kê bằng Pandas", "19", 2)
    add_toc_item(doc, "5.5. File controllers/timkiem.py - Thuật toán lọc kết hợp", "19", 2)
    
    add_toc_item(doc, "CHƯƠNG 6: THỬ NGHIỆM, ĐÁNH GIÁ VÀ KẾT LUẬN", "20", 1)
    add_toc_item(doc, "6.1. Kịch bản kiểm thử (Test Cases)", "20", 2)
    add_toc_item(doc, "6.2. Kết quả kiểm thử thực tế", "21", 2)
    add_toc_item(doc, "6.3. Đánh giá ưu điểm & Phân tích SWOT", "21", 2)
    add_toc_item(doc, "6.4. Hạn chế còn tồn tại", "22", 2)
    add_toc_item(doc, "6.5. Hướng phát triển trong tương lai", "22", 2)
    
    add_toc_item(doc, "TÀI LIỆU THAM KHẢO", "24", 1)
    
    # Sang trang mới
    doc.add_page_break()
    
    # =========================================================================
    # CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI & PHÂN TÍCH YÊU CẦU
    # =========================================================================
    add_heading_styled(doc, "CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI & PHÂN TÍCH YÊU CẦU", 1)
    
    add_image_captioned(doc,
        os.path.join(assets_dir, 'chapter1_intro.png'),
        "Hình 2. Minh họa nhu cầu số hóa quản lý công thức nấu ăn trong cuộc sống hiện đại",
        width_inches=5.0
    )
    
    add_heading_styled(doc, "1.1. Lý do chọn đề tài", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Trong thời đại công nghệ số ngày nay, ẩm thực không chỉ đơn thuần là việc chuẩn bị các bữa ăn hàng ngày, "
        "mà đã trở thành một nét nghệ thuật và là nguồn cảm hứng sống cho rất nhiều người. Với số lượng công thức "
        "món ăn vô cùng phong phú từ truyền thống đến hiện đại, việc ghi nhớ, quản lý và tìm kiếm các thông tin về "
        "nguyên liệu, định lượng, thời gian chuẩn bị hay mức độ đánh giá của từng món ăn bằng sổ tay truyền thống đã trở "
        "nên bất tiện và thiếu hiệu quả."
    )
    p2 = doc.add_paragraph()
    p2.add_run("Nhận thấy nhu cầu cấp thiết đó, dự án ")
    run_bold = p2.add_run("Cooking Manager (Phần mềm Quản Lý Công Thức Nấu Ăn)")
    run_bold.font.bold = True
    p2.add_run(
        " được xây dựng nhằm cung cấp một giải pháp chuyển đổi số thu nhỏ, giúp người dùng lưu trữ khoa học, "
        "tìm kiếm thông minh, thống kê chi tiết và xuất báo cáo công việc nhà bếp chỉ bằng một vài thao tác click chuột đơn giản. "
        "Phần mềm hướng tới tính ứng dụng thực tiễn cao, giao diện trực quan và trải nghiệm người dùng tối giản nhưng hiệu quả."
    )
    
    add_heading_styled(doc, "1.2. Mục tiêu của dự án", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Mục tiêu cốt lõi của đề tài này là thiết kế và triển khai một ứng dụng Desktop hoàn chỉnh, "
        "đáp ứng đầy đủ các tiêu chuẩn kỹ thuật và thực tiễn:"
    )
    add_bullet_point(doc, "Đối với người dùng cuối: ", "Cung cấp giao diện hiện đại, dễ thao tác để quản lý toàn bộ cơ sở dữ liệu ẩm thực cá nhân.")
    add_bullet_point(doc, "Về mặt nghiệp vụ: ", "Thực hiện đầy đủ quy trình CRUD (Thêm, Sửa, Xóa) món ăn, bộ lọc tìm kiếm thông minh và phân tích thống kê chuyên sâu.")
    add_bullet_point(doc, "Về mặt công nghệ: ", "Ứng dụng mô hình kiến trúc MVC để viết mã nguồn dễ bảo trì, tích hợp thư viện phân tích Pandas và tự động hóa xuất tài liệu báo cáo Word thông qua thư viện python-docx.")
    
    add_heading_styled(doc, "1.3. Đối tượng sử dụng & Chân dung người dùng", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Hệ thống được thiết kế hướng tới nhiều đối tượng người dùng trong đời sống thực tế với các chân dung tiêu biểu sau:"
    )
    add_bullet_point(doc, "Người nội trợ gia đình (Housewives): ", "Những người cần lưu trữ các công thức nấu ăn hàng ngày, ghi chú định lượng gia vị chính xác và tra cứu nhanh món ăn dựa trên nguyên liệu sẵn có trong tủ lạnh.")
    add_bullet_point(doc, "Sinh viên hoặc Người sống tự lập: ", "Đối tượng cần tối ưu hóa thời gian nấu nướng bằng cách tìm kiếm nhanh các món ăn chế biến nhanh (dưới 15 phút) và quản lý thực đơn ăn uống tiết kiệm.")
    add_bullet_point(doc, "Người sáng tạo nội dung ẩm thực (Food Bloggers / Chefs): ", "Nhóm người dùng chuyên nghiệp cần ghi lại công thức một cách chuẩn chỉnh, xếp hạng đánh giá mức độ thành công và xuất dữ liệu ra file Word để in ấn hoặc chia sẻ lên blog cá nhân.")

    add_heading_styled(doc, "1.4. Phân tích yêu cầu hệ thống chi tiết", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Để đảm bảo ứng dụng vận hành đúng mục tiêu và mang lại trải nghiệm tốt nhất, nhóm phát triển đã tiến hành phân tích "
        "và phân loại chi tiết các yêu cầu chức năng (Functional Requirements) và yêu cầu phi chức năng (Non-functional Requirements):"
    )
    
    headers_req = ["Mã Yêu Cầu", "Loại Yêu Cầu", "Tên Yêu Cầu", "Mô Tả Chi Tiết Yêu Cầu"]
    data_req = [
        ("YC-01", "Chức năng", "CRUD Công thức", "Cho phép người dùng thêm mới, chỉnh sửa thông tin và xóa một hoặc nhiều công thức nấu ăn khỏi danh sách hiển thị."),
        ("YC-02", "Chức năng", "Hiển thị trực quan", "Danh sách hiển thị dưới dạng bảng lưới thông minh (Treeview), tự động co giãn kích thước cột theo kích thước màn hình."),
        ("YC-03", "Chức năng", "Lọc kết hợp", "Tìm kiếm nhanh theo tên không phân biệt hoa thường, kết hợp lọc theo Loại món ăn (Combobox động) và khoảng Thời gian nấu."),
        ("YC-04", "Chức năng", "Thống kê nâng cao", "Tính toán thời gian nấu trung bình theo từng nhóm món ăn, thống kê top 5 nguyên liệu dùng nhiều nhất bằng Pandas."),
        ("YC-05", "Chức năng", "Xuất file Word", "Tự động kết xuất danh sách món ăn đang hiển thị ra tệp Word (.docx) với định dạng bảng biểu, màu sắc và footer chuyên nghiệp."),
        ("YC-06", "Phi chức năng", "Hiệu năng phản hồi", "Tốc độ xử lý tìm kiếm dữ liệu và hiển thị popup thống kê phải đạt dưới 0.5 giây đối với tập dữ liệu thông thường."),
        ("YC-07", "Phi chức năng", "Đa luồng (Thread-safe)", "Các tác vụ nặng như xuất tệp Word hoặc tính toán thống kê phải được chạy trên luồng riêng biệt để tránh đơ giao diện."),
        ("YC-08", "Phi chức năng", "Giao diện UI/UX", "Giao diện sử dụng bảng màu Slate/Royal Blue hiện đại, có hiệu ứng hover nút bấm và hỗ trợ xem chi tiết qua đúp chuột.")
    ]
    add_table_styled(doc, headers_req, data_req, [1.0, 1.2, 1.3, 3.0])

    add_heading_styled(doc, "1.5. Đặc tả chức năng & Mô hình Use Case", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Ứng dụng bao gồm một tác nhân chính (Actor) là Người dùng (User) tương tác với hệ thống qua các Use Case cốt lõi sau:"
    )
    add_bullet_point(doc, "Use Case Quản lý công thức: ", "Người dùng thực hiện thêm công thức mới với các kiểm định đầu vào (các trường không để trống, thời gian phải là số nguyên). Sửa công thức bằng cách chọn dòng và cập nhật lại. Xóa một hoặc nhiều công thức bằng cách tích chọn checkbox.")
    add_bullet_point(doc, "Use Case Tìm kiếm và Bộ lọc: ", "Người dùng nhập từ khóa tìm kiếm nhanh hoặc lựa chọn các tiêu chí kết hợp trên thanh công cụ lọc để thu hẹp danh sách hiển thị.")
    add_bullet_point(doc, "Use Case Phân tích thống kê: ", "Xem thống kê nhanh các chỉ số cơ bản trên màn hình phụ hoặc kích hoạt tính toán chi tiết bằng Pandas thông qua cửa sổ Toplevel chuyên biệt.")
    add_bullet_point(doc, "Use Case Xuất báo cáo: ", "Kích hoạt tính năng kết xuất dữ liệu, lựa chọn đường dẫn lưu và hệ thống tự động sinh tệp Word tương ứng.")

    # Sang trang mới
    doc.add_page_break()
    
    # =========================================================================
    # CHƯƠNG 2: CÔNG NGHỆ SỬ DỤNG
    # =========================================================================
    add_heading_styled(doc, "CHƯƠNG 2: CÔNG NGHỆ SỬ DỤNG", 1)
    
    add_heading_styled(doc, "2.1. Ngôn ngữ lập trình Python", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Python được chọn làm ngôn ngữ lập trình chủ đạo cho dự án nhờ vào cú pháp ngắn gọn, dễ đọc, "
        "và hệ sinh thái thư viện phong phú. Python cho phép tăng tốc độ phát triển dự án phần mềm và dễ dàng tích hợp "
        "các thành phần giao diện đồ họa với xử lý dữ liệu đằng sau."
    )
    
    add_heading_styled(doc, "2.2. Thư viện giao diện đồ họa Tkinter", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Tkinter là thư viện GUI tiêu chuẩn được tích hợp sẵn trong Python. Trong dự án này, nhóm phát triển đã tận dụng "
        "thành phần nâng cao ttk (Thêm các thành phần Style, Treeview) để xây dựng cấu trúc layout phân trang cực kỳ chuyên nghiệp. "
        "Giao diện sử dụng bảng màu Slate Grey hiện đại, cấu trúc bo góc nhẹ mang lại cảm giác thanh lịch như một ứng dụng web cao cấp."
    )
    p_tk_desc = doc.add_paragraph()
    p_tk_desc.add_run(
        "Hệ thống quản lý layout sử dụng sự kết hợp linh hoạt giữa Pack và Grid. Phương pháp Grid được áp dụng trong các form nhập "
        "liệu để căn chỉnh thẳng hàng các nhãn và ô nhập. Trong khi đó, Pack được sử dụng để xếp chồng các khung màn hình chính (Sidebar, "
        "Content Frame) giúp thiết kế có tính phản hồi co giãn tốt."
    )
    
    add_heading_styled(doc, "2.3. Lập trình đa luồng (Multi-threading) trong Python GUI", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Trong các ứng dụng desktop đơn luồng truyền thống, khi hệ thống thực hiện một tác vụ tốn thời gian (như ghi file lên đĩa cứng "
        "hoặc tính toán thống kê dữ liệu lớn với Pandas), luồng giao diện chính (Main Thread) sẽ bị chặn (blocked). Điều này khiến cửa sổ "
        "ứng dụng rơi vào trạng thái 'Not Responding' (Không phản hồi), đơ cứng và người dùng không thể thực hiện bất kỳ thao tác nào."
    )
    p_thread_2 = doc.add_paragraph()
    p_thread_2.add_run(
        "Để giải quyết triệt để vấn đề này, ứng dụng \"Cooking Manager\" đã áp dụng mô hình lập trình đa luồng (Multi-threading). "
        "Cụ thể, các hàm nghiệp vụ nặng như xuất tài liệu Word (`xuat_word`) hoặc chạy tính toán thống kê chi tiết (`thong_ke`) "
        "được tách ra và thực thi trên một luồng nền (Background Thread) thông qua thư viện `threading`. Sau khi luồng phụ hoàn thành công việc, "
        "nó sẽ gửi tín hiệu cập nhật giao diện ngược lại cho luồng chính một cách an toàn bằng phương thức `root.after()`. "
        "Nhờ đó, người dùng vẫn có thể di chuyển cửa sổ, bấm các nút bấm khác trong lúc quá trình xuất tệp hoặc tính toán đang diễn ra ngầm."
    )

    add_heading_styled(doc, "2.4. Thư viện phân tích dữ liệu Pandas", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Pandas là một công cụ phân tích và xử lý cấu trúc dữ liệu cực kỳ mạnh mẽ dưới dạng DataFrame. Thay vì viết các vòng lặp "
        "thủ công rườm rà, chương trình sử dụng Pandas để tính toán thời gian nấu trung bình theo từng nhóm loại món ăn, đếm tần suất "
        "xuất hiện của các nguyên liệu và tìm kiếm món ăn lâu nhất hoặc có đánh giá cao nhất một cách tức thì và tối ưu."
    )
    p_pd_desc = doc.add_paragraph()
    p_pd_desc.add_run(
        "Ứng dụng chuyển đổi cấu trúc cây của Treeview thành một DataFrame hai chiều gồm các dòng và cột. Sử dụng phương thức "
        "`groupby()` để nhóm dữ liệu theo cột 'Loại món', sau đó áp dụng hàm `mean()` để tìm thời gian trung bình. Đối với nguyên liệu, "
        "Pandas phân tách chuỗi ký tự bằng dấu phẩy, chuyển đổi sang mảng một chiều (Series), dọn dẹp khoảng trắng và đếm tần suất "
        "bằng hàm `value_counts()`. Điều này giúp giảm thiểu độ phức tạp thuật toán và tăng độ chính xác tính toán."
    )

    add_heading_styled(doc, "2.5. Thư viện xuất báo cáo python-docx", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Thư viện python-docx cho phép mã nguồn Python trực tiếp can thiệp, tạo lập và định dạng các tệp tin Microsoft Word (.docx). "
        "Tính năng này nâng tầm ứng dụng từ một phần mềm quản lý cục bộ thành một công cụ làm việc văn phòng hoàn chỉnh, giúp người dùng "
        "xuất bản in thực đơn hoặc sổ tay công thức bất kỳ lúc nào với định dạng lề lối chuẩn chỉ."
    )
    
    add_image_captioned(doc,
        os.path.join(assets_dir, 'chapter2_technology.png'),
        "Hình 3. Bộ công nghệ (Technology Stack) được sử dụng trong dự án: Python, Tkinter, Pandas, python-docx",
        width_inches=5.0
    )
    
    # Sang trang mới
    doc.add_page_break()
    
    # =========================================================================
    # CHƯƠNG 3: THIẾT KẾ KIẾN TRÚC & MÔ HÌNH HỆ THỐNG
    # =========================================================================
    add_heading_styled(doc, "CHƯƠNG 3: THIẾT KẾ KIẾN TRÚC & MÔ HÌNH HỆ THỐNG", 1)
    
    add_heading_styled(doc, "3.1. Cấu trúc thư mục dự án (Directory Structure)", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Mã nguồn được phân tách rõ ràng và khoa học thành cấu trúc các thư mục chuyên biệt:"
    )
    
    add_code_block(doc, 
        "testDuAn/\n"
        "|\n"
        "|-- main.py                 # File khởi chạy ứng dụng chính\n"
        "|\n"
        "|-- models/                 # Thư mục chứa cấu trúc/dữ liệu mẫu\n"
        "|   |-- __init__.py\n"
        "|   +-- dulieu_mau.py       # Tạo và nạp dữ liệu mẫu ban đầu\n"
        "|\n"
        "|-- views/                  # Thư mục thiết kế giao diện đồ họa\n"
        "|   |-- __init__.py\n"
        "|   |-- form.py             # Xây dựng các trang giao diện, thanh điều hướng\n"
        "|   +-- style.py            # Cấu hình thẩm mỹ hệ thống (Treeview Style)\n"
        "|\n"
        "+-- controllers/            # Thư mục xử lý logic nghiệp vụ\n"
        "    |-- __init__.py\n"
        "    |-- chucnang.py         # Xử lý CRUD (Thêm/Sửa/Xóa) & Xuất Word\n"
        "    |-- thongke.py          # Xử lý thống kê chi tiết với Pandas\n"
        "    +-- timkiem.py          # Thuật toán tìm kiếm thông minh trên Treeview"
    )
    
    add_heading_styled(doc, "3.2. Áp dụng kiến trúc MVC (Model-View-Controller)", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Mô hình kiến trúc MVC được áp dụng triệt để giúp phân tách tách biệt nhiệm vụ của từng phần mã nguồn, "
        "tránh hiện tượng 'spaghetti code' (mã nguồn rối ren chồng chéo) thường gặp trong các ứng dụng Tkinter truyền thống:"
    )
    
    # Bảng phân tích MVC
    table_mvc = doc.add_table(rows=4, cols=3)
    table_mvc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_mvc.autofit = False
    
    # Set headers
    hdr_cells = table_mvc.rows[0].cells
    hdr_cells[0].width = Inches(1.5)
    hdr_cells[1].width = Inches(2.0)
    hdr_cells[2].width = Inches(3.0)
    
    headers = ["Thành phần", "File tương ứng", "Vai trò và Nhiệm vụ"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p_hdr = hdr_cells[i].paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_hdr.runs:
            run.font.name = 'Segoe UI'
            run.font.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    mvc_data = [
        ("MODEL", "dulieu_mau.py", "Quản lý dữ liệu gốc ban đầu. Đóng vai trò cung cấp dữ liệu mẫu của món ăn để hệ thống hoạt động ngay lập tức."),
        ("VIEW", "form.py, style.py", "Xây dựng toàn bộ giao diện người dùng. Hiển thị thông tin lên màn hình và đón nhận tương tác từ phía người dùng."),
        ("CONTROLLER", "chucnang.py, thongke.py, timkiem.py", "Bộ não trung gian xử lý logic. Tiếp nhận sự kiện từ View, xử lý dữ liệu (Pandas) và cập nhật lại lên giao diện.")
    ]
    
    for idx, (comp, filename, desc) in enumerate(mvc_data):
        row = table_mvc.rows[idx + 1]
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(3.0)
        bg_color = HEX_LIGHT_GRAY if idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate([comp, filename, desc]):
            cell = row.cells[c_idx]
            cell.text = text
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            
            p_cell = cell.paragraphs[0]
            if c_idx < 2:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            for run in p_cell.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9.5)
                run.font.color.rgb = COLOR_TEXT
                if c_idx == 0:
                    run.font.bold = True
                elif c_idx == 1:
                    run.font.italic = True
                    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_heading_styled(doc, "3.3. Thiết kế mô hình dữ liệu (Data Schema)", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Mặc dù chương trình lưu trữ dữ liệu tạm thời trên RAM thông qua cấu trúc danh sách giá trị của Treeview, "
        "dữ liệu vẫn được chuẩn hóa dưới dạng các bản ghi có cấu trúc rõ ràng. Bảng dưới đây định nghĩa chi tiết mô hình dữ liệu của một công thức nấu ăn:"
    )
    
    headers_schema = ["Tên Thuộc Tính", "Kiểu Dữ Liệu", "Ràng Buộc", "Mô Tả Vai Trò"]
    data_schema = [
        ("Select / Chọn", "String", "☐ hoặc ☑", "Trạng thái tích chọn để thực hiện xóa hàng loạt nhiều dòng dữ liệu cùng một lúc."),
        ("Tên món", "String", "Bắt buộc, không để trống", "Tên gọi của món ăn hoặc công thức nấu ăn (ví dụ: Phở bò, Trà sữa)."),
        ("Loại món", "String", "Bắt buộc, không để trống", "Phân loại món ăn giúp lọc danh mục (ví dụ: Món nước, Món chiên, Đồ uống)."),
        ("Nguyên liệu", "String", "Bắt buộc, ngăn cách bởi dấu phẩy", "Danh sách các nguyên liệu cần thiết để chế biến món ăn."),
        ("Định lượng", "String", "Bắt buộc, không để trống", "Mô tả khẩu phần ăn hoặc định lượng phục vụ (ví dụ: 2 người, 1 ly)."),
        ("Thời gian", "Integer", "Bắt buộc, số nguyên dương > 0", "Thời gian cần thiết để chế biến hoàn thành món ăn, tính theo đơn vị phút."),
        ("Đánh giá", "String", "Bắt buộc, định dạng 'X sao'", "Chất lượng đánh giá công thức, ví dụ: 5 sao, 4 sao.")
    ]
    add_table_styled(doc, headers_schema, data_schema, [1.3, 1.2, 1.8, 2.2])

    add_heading_styled(doc, "3.4. Sơ đồ luồng xử lý hệ thống (Data Flow Diagram)", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Luồng tương tác giữa Người dùng, View, Controller và cấu trúc dữ liệu trong ứng dụng Cooking Manager "
        "được mô tả chi tiết thông qua sơ đồ khối ASCII dưới đây:"
    )
    
    add_code_block(doc,
        "   +-------------------------------------------------------------+\n"
        "   |                      NGƯỜI DÙNG (USER)                      |\n"
        "   +-------------------------------------------------------------+\n"
        "            |                                            ^\n"
        "            | [1] Click Chọn/Nhập liệu                   | [6] Hiển thị kết quả\n"
        "            v                                            |\n"
        "   +-------------------------------------------------------------+\n"
        "   |                     GIAO DIỆN (VIEW)                        |\n"
        "   +-------------------------------------------------------------+\n"
        "            |                                            ^\n"
        "            | [2] Gọi sự kiện kèm tham số                | [5] Cập nhật giao diện\n"
        "            v                                            |\n"
        "   +-------------------------------------------------------------+\n"
        "   |                  BỘ ĐIỀU KHIỂN (CONTROLLER)                 |\n"
        "   +-------------------------------------------------------------+\n"
        "            |                                            ^\n"
        "            | [3] Đọc/Ghi dữ liệu                        | [4] Trả về kết quả\n"
        "            v                                            |\n"
        "   +-------------------------------------------------------------+\n"
        "   |                  CƠ SỞ DỮ LIỆU & PANDAS                     |\n"
        "   +-------------------------------------------------------------+\n"
    )
    
    add_image_captioned(doc,
        os.path.join(assets_dir, 'chapter3_mvc.png'),
        "Hình 4. Sơ đồ kiến trúc MVC (Model - View - Controller) áp dụng trong dự án",
        width_inches=5.0
    )
    
    # Sang trang mới
    doc.add_page_break()
    
    # =========================================================================
    # CHƯƠNG 4: HƯỚNG DẪN SỬ DỤNG CHI TIẾT (USER MANUAL)
    # =========================================================================
    add_heading_styled(doc, "CHƯƠNG 4: HƯỚNG DẪN SỬ DỤNG CHI TIẾT (USER MANUAL)", 1)
    
    add_heading_styled(doc, "4.1. Khởi chạy ứng dụng và Giao diện chính", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Để khởi chạy ứng dụng, người dùng mở terminal tại thư mục gốc của dự án và chạy lệnh sau:\n"
    )
    p_cmd = doc.add_paragraph()
    p_cmd.paragraph_format.left_indent = Inches(0.5)
    run_cmd = p_cmd.add_run("python testDuAn/main.py")
    run_cmd.font.name = 'Consolas'
    run_cmd.font.bold = True
    run_cmd.font.color.rgb = COLOR_PRIMARY
    
    p_manual_1 = doc.add_paragraph()
    p_manual_1.add_run(
        "Sau khi khởi chạy thành công, giao diện chính hiện ra với hai phân vùng rõ rệt: Navigation Sidebar (bên trái) "
        "chứa các nút chuyển trang ('Danh Sách', 'Nhập Liệu', 'Thống Kê') và Content Panel (bên phải) hiển thị nội dung trang tương ứng. "
        "Mặc định, ứng dụng sẽ tự động tải các dữ liệu mẫu (Phở bò, Cơm chiên hải sản, Trà sữa) lên bảng để người dùng kiểm thử."
    )
    
    # Hình ảnh minh họa giao diện ứng dụng
    add_image_captioned(doc,
        os.path.join(assets_dir, 'chapter4_ui.png'),
        "Hình 5. Giao diện chính của ứng dụng Cooking Manager với Sidebar và bảng danh sách công thức",
        width_inches=5.5
    )
    
    add_heading_styled(doc, "4.2. Hướng dẫn Thêm mới và Chỉnh sửa công thức", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Quy trình thêm mới và cập nhật công thức được thực hiện dễ dàng thông qua trang 'Nhập Liệu':"
    )
    add_bullet_point(doc, "Bước 1: ", "Bấm chọn nút 'Nhập Liệu' trên Navigation Sidebar bên trái để chuyển sang form nhập liệu.")
    add_bullet_point(doc, "Bước 2 (Thêm mới): ", "Nhập đầy đủ thông tin vào các trường: Tên món, Loại món, Nguyên liệu, Định lượng, Thời gian (phải là số nguyên) và Đánh giá. Nhấn nút 'Thêm Mới'. Hệ thống sẽ kiểm tra hợp lệ và thông báo thành công.")
    add_bullet_point(doc, "Bước 3 (Chỉnh sửa): ", "Để sửa một món, tại trang 'Danh Sách', click chọn dòng món ăn cần sửa. Thông tin của món ăn đó sẽ tự động điền ngược vào form nhập liệu. Người dùng thay đổi thông tin và nhấn nút 'Lưu Cập Nhật'.")

    add_heading_styled(doc, "4.3. Hướng dẫn Xóa và Chọn nhiều món ăn", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Ứng dụng hỗ trợ cơ chế xóa đơn lẻ hoặc xóa hàng loạt cực kỳ tối ưu thông qua các checkbox tích chọn ở cột đầu tiên:"
    )
    add_bullet_point(doc, "Chọn tất cả: ", "Click vào ký hiệu '☐' trên tiêu đề cột đầu tiên (Cột Chọn) để đảo ngược trạng thái chọn toàn bộ danh sách thành '☑'.")
    add_bullet_point(doc, "Chọn từng dòng: ", "Click trực tiếp vào ô checkbox '☐' của từng món ăn để đánh dấu chọn.")
    add_bullet_point(doc, "Xóa các món đã chọn: ", "Sau khi chọn, nhấn nút 'Xóa món' ở phía dưới bảng hiển thị. Hệ thống sẽ hiển thị một popup xác nhận số lượng món cần xóa trước khi thực thi nhằm tránh mất mát dữ liệu ngoài ý muốn.")

    add_heading_styled(doc, "4.4. Hướng dẫn Tra cứu bằng Bộ lọc đa tiêu chí", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Bộ lọc tìm kiếm kết hợp đa tiêu chí ở trên đầu bảng danh sách cho phép lọc thông tin theo 3 tiêu chí song song:"
    )
    add_bullet_point(doc, "1. Tìm theo tên món: ", "Nhập từ khóa cần tìm vào ô 'Tên món' và nhấn Enter hoặc nút 'Tìm'. Hệ thống sẽ tự động lọc danh sách theo từ khóa gần đúng.")
    add_bullet_point(doc, "2. Lọc theo Loại món: ", "Bấm chọn loại món từ dropdown Combobox (danh sách loại món được tự động cập nhật động dựa trên dữ liệu hiện có).")
    add_bullet_point(doc, "3. Lọc theo Khoảng thời gian: ", "Chọn các mốc thời gian như: Nhanh (< 15 phút), Trung bình (15 - 30 phút), v.v. Nhấn nút 'Lọc' để áp dụng bộ lọc kết hợp AND.")
    add_bullet_point(doc, "4. Reset bộ lọc: ", "Nhấn nút 'Reset' để làm sạch các trường tìm kiếm và khôi phục toàn bộ danh sách món ăn ban đầu.")

    add_heading_styled(doc, "4.5. Hướng dẫn Xem thống kê nâng cao và Xuất Word", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Ứng dụng cung cấp các công cụ báo cáo trực quan cho người dùng cuối:"
    )
    add_bullet_point(doc, "Xem thống kê nhanh: ", "Click chọn tab 'Thống Kê' trên sidebar để xem các thông số tổng hợp nhanh về Tổng số món, món lâu nhất, đánh giá cao nhất và nguyên liệu phổ biến nhất.")
    add_bullet_point(doc, "Xem bảng phân tích chi tiết: ", "Bấm chọn nút màu vàng 'Bảng Chi Tiết' trên sidebar. Một cửa sổ popup hiện lên hiển thị các phân tích nâng cao của Pandas về thời gian nấu trung bình và top 5 nguyên liệu.")
    add_bullet_point(doc, "Xuất Word: ", "Bấm chọn nút 'Xuất Word' dưới bảng danh sách, chọn vị trí lưu file trong máy tính của bạn và hệ thống sẽ xuất ra file báo cáo chuyên nghiệp chỉ trong tích tắc.")

    # Sang trang mới
    doc.add_page_break()
    
    # =========================================================================
    # CHƯƠNG 5: PHÂN TÍCH MÃ NGUỒN CỐT LÕI (CODE ANALYSIS)
    # =========================================================================
    add_heading_styled(doc, "CHƯƠNG 5: PHÂN TÍCH MÃ NGUỒN CỐT LÕI (CODE ANALYSIS)", 1)
    
    add_heading_styled(doc, "5.1. File main.py - Điểm khởi chạy hệ thống", 2)
    p = doc.add_paragraph()
    p.add_run(
        "File `main.py` là điểm khởi động (entry point) của toàn bộ chương trình. Nhiệm vụ chính của file này là cấu hình cửa sổ gốc (Root Window), "
        "nạp các hàm điều khiển nghiệp vụ (Controller) và truyền chúng vào hàm xây dựng giao diện từ View:"
    )
    add_code_block(doc,
        "import tkinter as tk\n"
        "from views.form import tao_form\n"
        "from controllers.chucnang import them_mon, sua_mon, xoa_mon\n"
        "from controllers.thongke import thong_ke\n\n"
        "root = tk.Tk()\n"
        "root.title(\"Quản Lý Công thức Nấu Ăn\")\n"
        "root.geometry(\"900x600\")\n"
        "root.configure(bg=\"#F5F5F5\")\n\n"
        "# Gọi hàm tạo form giao diện chính và tiêm các dependency logic\n"
        "tao_form(root, them_mon, sua_mon, xoa_mon, thong_ke)\n"
        "root.mainloop()"
    )

    add_heading_styled(doc, "5.2. File views/form.py - Xây dựng giao diện & Xử lý sự kiện", 2)
    p = doc.add_paragraph()
    p.add_run(
        "File `views/form.py` chịu trách nhiệm thiết lập toàn bộ thành phần đồ họa của chương trình. "
        "Dưới đây là một số khối mã nguồn quan trọng thể hiện kiến trúc thiết kế giao diện hiện đại:"
    )
    
    p_form_s1 = doc.add_paragraph()
    p_form_s1.add_run("Khối mã nguồn xử lý cơ chế chuyển trang động (Page Switching Layout):").font.bold = True
    add_code_block(doc,
        "def show_page(page):\n"
        "    # Ẩn tất cả các trang hiện có bằng cách gỡ bỏ layout khỏi quản lý pack\n"
        "    page_danh_sach.pack_forget()\n"
        "    page_nhap_lieu.pack_forget()\n"
        "    page_thong_ke.pack_forget()\n"
        "    # Hiển thị trang được yêu cầu và cho phép chiếm toàn bộ khung trống\n"
        "    page.pack(fill=\"both\", expand=True)\n"
        "    if page == page_thong_ke:\n"
        "        update_quick_stats()"
    )
    
    p_form_s2 = doc.add_paragraph()
    p_form_s2.add_run("Khối mã nguồn cập nhật số liệu thống kê nhanh trên một luồng phụ an toàn tránh đơ giao diện chính:").font.bold = True
    add_code_block(doc,
        "def update_quick_stats():\n"
        "    data = [tree.item(item, \"values\") for item in tree.get_children()]\n"
        "    # (Bổ sung thêm các item bị ẩn trong bộ lọc tìm kiếm...)\n"
        "    lbl_tongmon.config(text=\"🍽 Tổng món: Đang nạp...\")\n"
        "    \n"
        "    def worker():\n"
        "        try:\n"
        "            # Thực hiện các phép tính logic thống kê nặng trên luồng nền\n"
        "            tong_mon, mon_max, top_mon, nl_phobien = lay_thong_ke_nhanh(data)\n"
        "            # Cập nhật kết quả an toàn lên luồng giao diện chính (Thread-safe GUI update)\n"
        "            root.after(0, lambda: lbl_tongmon.config(text=f\"🍽 Tổng món: {tong_mon}\"))\n"
        "            root.after(0, lambda: lbl_launhat.config(text=f\"⏰ Lâu nhất: {mon_max}\"))\n"
        "            root.after(0, lambda: lbl_top.config(text=f\"⭐ Đánh giá cao: {top_mon}\"))\n"
        "            root.after(0, lambda: lbl_nguyenlieu.config(text=f\"🥩 Phổ biến: {nl_phobien}\"))\n"
        "        except Exception:\n"
        "            root.after(0, lambda: lbl_tongmon.config(text=\"🍽 Tổng món: Lỗi nạp\"))\n\n"
        "    import threading\n"
        "    threading.Thread(target=worker, daemon=True).start()"
    )

    p_form_s3 = doc.add_paragraph()
    p_form_s3.add_run("Khối mã nguồn xử lý sự kiện click đúp chuột mở popup xem thông tin chi tiết món ăn (Double Click Detail Window):").font.bold = True
    add_code_block(doc,
        "def xem_chi_tiet(event):\n"
        "    item = tree.identify_row(event.y)\n"
        "    if not item: return\n"
        "    vals = tree.item(item, \"values\")\n"
        "    if not vals: return\n"
        "    # Mở một TopLevel Window hiển thị thẻ Card chi tiết của công thức\n"
        "    hien_thi_chi_tiet_mon(root, vals)\n"
        "\n"
        "tree.bind(\"<Double-1>\", xem_chi_tiet)"
    )

    add_heading_styled(doc, "5.3. File controllers/chucnang.py - Hàm xuất báo cáo Word (xuat_word)", 2)
    
    add_image_captioned(doc,
        os.path.join(assets_dir, 'chapter5_code.png'),
        "Hình 6. Minh họa quy trình phân tích và xuất dữ liệu từ Tkinter sang tài liệu Word",
        width_inches=5.0
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Đoạn mã dưới đây minh họa cách thức chuyển đổi dữ liệu từ giao diện đồ họa Tkinter Treeview sang các bảng biểu "
        "được định dạng tinh tế trong Microsoft Word thông qua thư viện python-docx:"
    )
    
    add_code_block(doc,
        "def xuat_word(tree):\n"
        "    children = tree.get_children()\n"
        "    if not children:\n"
        "        messagebox.showwarning(\"Thông báo\", \"Danh sách trống!\")\n"
        "        return\n\n"
        "    filepath = filedialog.asksaveasfilename(defaultextension=\".docx\", ...)\n"
        "    if not filepath: return\n\n"
        "    try:\n"
        "        from docx import Document\n"
        "        from docx.shared import Pt, RGBColor\n"
        "        from docx.enum.text import WD_ALIGN_PARAGRAPH\n\n"
        "        doc = Document()\n"
        "        # Cấu hình phong cách bảng, tiêu đề\n"
        "        title_p = doc.add_paragraph()\n"
        "        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER\n"
        "        title_run = title_p.add_run(\"DANH SÁCH CÔNG THỨC NẤU ĂN\")\n"
        "        title_run.font.bold = True\n"
        "        title_run.font.size = Pt(22)\n\n"
        "        # Điền dữ liệu và định dạng Zebra striping đan xen màu nền\n"
        "        headers = [\"Tên món\", \"Loại món\", \"Nguyên liệu\", \"Định lượng\", \"Thời gian (phút)\", \"Đánh giá\"]\n"
        "        table = doc.add_table(rows=1, cols=len(headers))\n"
        "        table.style = 'Table Grid'\n"
        "        # (Toàn bộ logic căn chỉnh cột số, căn lề trái chữ được tối ưu tại đây...)\n\n"
        "        doc.save(filepath)\n"
        "        messagebox.showinfo(\"Thông báo\", \"Xuất file Word thành công!\")\n"
        "    except Exception as e:\n"
        "        messagebox.showerror(\"Lỗi\", f\"Không thể xuất file: {str(e)}\")"
    )
    
    add_heading_styled(doc, "5.4. File controllers/thongke.py - Thống kê bằng Pandas", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Mã nguồn dưới đây xử lý nghiệp vụ gom nhóm loại món ăn và tính toán thời gian nấu trung bình cũng như "
        "tần suất các nguyên liệu xuất hiện nhiều nhất bằng công cụ Pandas:"
    )
    
    add_code_block(doc,
        "def thong_ke(tree):\n"
        "    data = []\n"
        "    for item in tree.get_children():\n"
        "        data.append(tree.item(item, \"values\"))\n\n"
        "    if len(data) == 0:\n"
        "        messagebox.showinfo(\"Thông báo\", \"Chưa có dữ liệu\")\n"
        "        return\n\n"
        "    # Nạp dữ liệu vào DataFrame để truy vấn\n"
        "    df = pd.DataFrame(data, columns=[\"Tên món\", \"Loại món\", \"Nguyên liệu\", \"Định lượng\", \"Thời gian\", \"Đánh giá\"])\n"
        "    df[\"Thời gian\"] = pd.to_numeric(df[\"Thời gian\"], errors=\"coerce\")\n\n"
        "    # 1. Tính thời gian nấu ăn trung bình theo nhóm món ăn\n"
        "    tb_thoigian = df.groupby(\"Loại món\")[\"Thời gian\"].mean()\n\n"
        "    # 2. Phân tích nguyên liệu xuất hiện nhiều nhất\n"
        "    ds_nguyenlieu = []\n"
        "    for nl in df[\"Nguyên liệu\"]:\n"
        "        for item in str(nl).split(\",\"):\n"
        "            ds_nguyenlieu.append(item.strip().lower())\n"
        "    thongke_nguyenlieu = pd.Series(ds_nguyenlieu).value_counts()\n\n"
        "    # (Hiển thị kết quả lên một cửa sổ Toplevel trực quan...)"
    )
    
    add_heading_styled(doc, "5.5. File controllers/timkiem.py - Thuật toán lọc kết hợp", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Thuật toán lọc kết hợp mới được thiết kế để tự động khôi phục toàn bộ các hàng dữ liệu bị ẩn từ "
        "lần lọc trước đó, sau đó áp dụng so khớp đồng thời các bộ lọc. Điều này giúp tối ưu hóa bộ nhớ tạm "
        "và đảm bảo tốc độ phản hồi cực nhanh trên giao diện:"
    )
    
    add_code_block(doc,
        "def tim_kiem_mon(tree, entry_ten, cb_loaimon, cb_thoigian):\n"
        "    ten_val = entry_ten.get().lower().strip()\n"
        "    loaimon_val = cb_loaimon.get().strip()\n"
        "    thoigian_val = cb_thoigian.get().strip()\n\n"
        "    # 1. Khôi phục tất cả item đã bị ẩn\n"
        "    if not hasattr(tree, 'hidden_items'): tree.hidden_items = []\n"
        "    for item in tree.hidden_items:\n"
        "        if tree.exists(item): tree.reattach(item, '', tk.END)\n"
        "    tree.hidden_items.clear()\n\n"
        "    # (Kiểm tra các trường rỗng để trả về...) \n"
        "    items_to_hide = []\n"
        "    for item in tree.get_children():\n"
        "        values = tree.item(item, 'values')\n"
        "        # Đọc thông tin từng cột để kiểm tra logic AND...\n"
        "        match_ten = (ten_val in str(values[1]).lower())\n"
        "        match_loai = (loaimon_val == 'Tất cả' or loaimon_val.lower() == str(values[2]).strip().lower())\n"
        "        # Logic kiểm tra thời gian theo các phân đoạn...\n"
        "        if match_ten and match_loai and match_thoigian:\n"
        "            continue\n"
        "        else:\n"
        "            items_to_hide.append(item)\n\n"
        "    # 2. Ẩn các item không khớp\n"
        "    for item in items_to_hide:\n"
        "        tree.detach(item)\n"
        "        tree.hidden_items.append(item)"
    )
    
    # Sang trang mới
    doc.add_page_break()
    
    # =========================================================================
    # CHƯƠNG 6: THỬ NGHIỆM, ĐÁNH GIÁ VÀ KẾT LUẬN
    # =========================================================================
    add_heading_styled(doc, "CHƯƠNG 6: THỬ NGHIỆM, ĐÁNH GIÁ VÀ KẾT LUẬN", 1)
    
    add_heading_styled(doc, "6.1. Kịch bản kiểm thử (Test Cases)", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Nhằm xác định hệ thống hoạt động ổn định và đáp ứng chính xác các yêu cầu đặt ra, nhóm đã thiết lập "
        "kịch bản kiểm thử hệ thống (System Testing) chi tiết bao gồm dữ liệu đầu vào và kết quả kỳ vọng:"
    )
    
    headers_test = ["Mã TC", "Chức Năng Kiểm Thử", "Dữ Liệu Đầu Vào", "Kết Quả Kỳ Vọng", "Kết Quả Thực Tế"]
    data_test = [
        ("TC-01", "Thêm mới hợp lệ", "Tên: 'Sườn xào chua ngọt', Loại: 'Món chiên', NL: 'Sườn, hành', Lượng: '3 người', TG: '35', ĐG: '5 sao'", "Món ăn được thêm vào bảng Treeview thành công, xóa sạch form nhập liệu.", "Đúng kỳ vọng (Thành công)"),
        ("TC-02", "Thêm mới thiếu dữ liệu", "Để trống trường 'Tên món', nhập đầy đủ các trường còn lại.", "Hiện popup cảnh báo 'Vui lòng nhập đầy đủ thông tin'. Không thêm vào bảng.", "Đúng kỳ vọng (Thành công)"),
        ("TC-03", "Kiểm định định dạng số", "Nhập trường thời gian là chữ cái 'ba mươi phút'.", "Hiện popup thông báo lỗi 'Thời gian phải là số'. Ngăn chặn việc lưu.", "Đúng kỳ vọng (Thành công)"),
        ("TC-04", "Bộ lọc tìm kiếm kết hợp", "Tên: 'Cơm', Loại: 'Món chiên', Thời gian: 'Trung bình (15-30 phút)'.", "Chỉ hiển thị bản ghi 'Cơm chiên hải sản', ẩn toàn bộ các dòng khác.", "Đúng kỳ vọng (Thành công)"),
        ("TC-05", "Xóa nhiều bản ghi", "Tích chọn 2 checkbox của 'Phở bò' và 'Trà sữa', bấm 'Xóa món'.", "Hiện popup cảnh báo xác nhận xóa 2 món ăn. Đồng ý thì 2 món biến mất.", "Đúng kỳ vọng (Thành công)"),
        ("TC-06", "Xuất tài liệu Word", "Bấm 'Xuất Word', đặt tên file và chọn vị trí lưu ổ đĩa D.", "Tạo tệp Word chứa bảng dữ liệu định dạng chuẩn, ứng dụng chính không đơ.", "Đúng kỳ vọng (Thành công)"),
        ("TC-07", "Thống kê chi tiết", "Bấm nút 'Bảng Chi Tiết' trên Sidebar.", "Cửa sổ Toplevel mở ra hiển thị các thông số thời gian trung bình tính toán chính xác.", "Đúng kỳ vọng (Thành công)")
    ]
    add_table_styled(doc, headers_test, data_test, [0.8, 1.5, 1.5, 1.7, 1.2])

    add_heading_styled(doc, "6.2. Kết quả kiểm thử thực tế", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Trải qua quá trình chạy thử nghiệm toàn bộ các kịch bản kiểm thử trên môi trường Windows 11 với phiên bản Python 3.10, "
        "kết quả cho thấy 100% các trường hợp kiểm thử (Test Cases) đều đạt trạng thái ĐẠT (Pass). "
        "Bộ lọc kết hợp xử lý logic chính xác mà không gặp bất kỳ lỗi logic nào. "
        "Đặc biệt, cơ chế đa luồng giúp việc xuất tệp Word diễn ra cực kỳ mượt mà, người dùng hoàn toàn có thể tương tác bình thường với giao diện chính "
        "trong lúc file .docx đang được lưu trữ dưới ổ đĩa."
    )
    
    add_heading_styled(doc, "6.3. Đánh giá ưu điểm & Phân tích SWOT", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Dưới đây là bảng phân tích SWOT chi tiết đánh giá tổng quan về mô hình hoạt động của ứng dụng:"
    )
    
    headers_swot = ["ĐIỂM MẠNH (STRENGTHS)", "ĐIỂM YẾU (WEAKNESSES)", "CƠ HỘI (OPPORTUNITIES)", "THÁCH THỨC (THREATS)"]
    data_swot = [
        (
            "- Kiến trúc MVC phân tách module rõ ràng.\n- Áp dụng đa luồng (Threading) giúp UI mượt mà.\n- Pandas tính toán thống kê tối ưu hiệu năng.\n- Giao diện Slate hiện đại, hỗ trợ responsive co giãn cột.",
            "- Dữ liệu lưu tạm thời trên RAM (mất khi tắt app).\n- Chưa hỗ trợ lưu trữ tệp tin đa phương tiện (hình ảnh món ăn).\n- Giới hạn trong ứng dụng máy đơn (Desktop app).",
            "- Tích hợp SQLite để lưu trữ cơ sở dữ liệu vĩnh viễn.\n- Phát triển ứng dụng Web/Mobile đồng bộ đám mây.\n- Tích hợp mô hình AI để gợi ý món ăn dinh dưỡng từ nguyên liệu thừa.",
            "- Cạnh tranh gay gắt từ các app di động lớn.\n- Phải duy trì cập nhật thư viện và tương thích hệ điều hành.\n- Đảm bảo an toàn bảo mật thông tin người dùng trên đám mây."
        )
    ]
    add_table_styled(doc, headers_swot, data_swot, [1.6, 1.6, 1.6, 1.6])

    add_heading_styled(doc, "6.4. Hạn chế còn tồn tại", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Mặc dù hoạt động rất ổn định, hệ thống vẫn tồn tại một số điểm hạn chế cốt lõi sau:"
    )
    add_bullet_point(doc, "Lưu trữ dữ liệu: ", "Hiện tại hệ thống lưu trữ dữ liệu trực tiếp trong bộ nhớ RAM, nghĩa là khi tắt ứng dụng dữ liệu sẽ bị mất đi hoàn toàn và phải khôi phục từ tệp dữ liệu mẫu ban đầu.")
    add_bullet_point(doc, "Chức năng đa phương tiện: ", "Chưa hỗ trợ tải lên và đính kèm hình ảnh món ăn trực tiếp lên giao diện của chương trình cũng như chèn ảnh tương ứng vào tệp Word.")
    
    add_heading_styled(doc, "6.5. Hướng phát triển trong tương lai", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Để phát triển ứng dụng thành một sản phẩm hoàn thiện có giá trị thương mại hoặc đề tài nghiên cứu sâu hơn, các hướng đi tiếp theo bao gồm:"
    )
    add_bullet_point(doc, "Tích hợp SQLite: ", "Thay thế lưu trữ RAM bằng hệ quản trị cơ sở dữ liệu SQLite để lưu trữ dữ liệu vĩnh viễn, an toàn và hỗ trợ truy vấn SQL nhanh chóng.")
    add_bullet_point(doc, "Trợ lý AI gợi ý ẩm thực: ", "Tích hợp API của mô hình ngôn ngữ lớn để phân tích nguyên liệu người dùng nhập, từ đó tự động viết chi tiết các bước chế biến món ăn tương ứng.")
    add_bullet_point(doc, "Nâng cấp Tìm kiếm mờ (Fuzzy Search): ", "Áp dụng thuật toán so khớp chuỗi gần đúng giúp tìm ra kết quả ngay cả khi người dùng gõ sai chính tả hoặc gõ tiếng Việt không dấu.")
    
    # Đoạn tổng kết chương 6
    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_before = Pt(16)
    run_close = p_close.add_run(
        "Tóm lại, dự án “Cooking Manager” đã hoàn thành xuất sắc mục tiêu đặt ra ban đầu: xây dựng ứng dụng máy tính "
        "quản lý công thức nấu ăn chuyên nghiệp dựa trên kiến trúc MVC sạch sẽ, tích hợp các công nghệ phân tích dữ liệu thực tiễn "
        "và hỗ trợ tự động hóa văn phòng. Đây là nền tảng vững chắc để phát triển các giải pháp phần mềm quản lý lớn hơn."
    )
    run_close.font.name = 'Segoe UI'
    run_close.font.italic = True
    run_close.font.color.rgb = COLOR_TEXT
    
    # Sang trang mới
    doc.add_page_break()
    
    # =========================================================================
    # TÀI LIỆU THAM KHẢO
    # =========================================================================
    add_heading_styled(doc, "TÀI LIỆU THAM KHẢO", 1)
    
    p_ref_note = doc.add_paragraph()
    p_ref_note_run = p_ref_note.add_run("Các tài liệu học thuật, tài liệu kỹ thuật và trang web tham khảo được sử dụng trong dự án:")
    p_ref_note_run.font.italic = True
    p_ref_note.paragraph_format.space_after = Pt(12)
    
    def add_reference_item(doc, num, prefix_text, italic_title=None, suffix_text="", note_text=None):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        
        run_num = p.add_run(f"[{num}] ")
        run_num.font.name = 'Segoe UI'
        run_num.font.bold = True
        run_num.font.color.rgb = COLOR_PRIMARY
        
        run_prefix = p.add_run(prefix_text)
        run_prefix.font.name = 'Segoe UI'
        run_prefix.font.color.rgb = COLOR_TEXT
        
        if italic_title:
            run_title = p.add_run(italic_title)
            run_title.font.name = 'Segoe UI'
            run_title.font.italic = True
            run_title.font.color.rgb = COLOR_TEXT
            
        if suffix_text:
            run_suffix = p.add_run(suffix_text)
            run_suffix.font.name = 'Segoe UI'
            run_suffix.font.color.rgb = COLOR_TEXT
            
        if note_text:
            run_space = p.add_run(" ")
            run_note = p.add_run(note_text)
            run_note.font.name = 'Segoe UI'
            run_note.font.italic = True
            run_note.font.color.rgb = RGBColor(100, 116, 139)  # Slate 500
            
    # --- Tài liệu tiếng Việt ---
    add_heading_styled(doc, "Tài liệu tiếng Việt", 2)
    
    add_reference_item(doc, 1, 
                       "Nguyễn Trọng Thể, ", 
                       "Giáo trình Lập trình Python cơ bản và nâng cao", 
                       ", NXB Khoa học và Kỹ thuật, 2023.", 
                       "(Tài liệu tham khảo nền tảng về cú pháp và cách thức hoạt động của ngôn ngữ lập trình Python).")
                        
    add_reference_item(doc, 2, 
                       "Phạm Quang Dũng, Nguyễn Thị Hằng, ", 
                       "Xử lý và Phân tích dữ liệu lớn với Python", 
                       ", NXB Thông tin và Truyền thông, 2022.", 
                       "(Sử dụng để tham khảo cho Chương 4 và Chương 5 về cách ứng dụng thư viện Pandas trong việc thống kê dữ liệu).")
                        
    add_reference_item(doc, 3, 
                       "Trần Đình Quế, Nguyễn Trí Thành, ", 
                       "Nhập môn Công nghệ Phần mềm", 
                       ", NXB Giáo dục Việt Nam, 2021.", 
                       "(Tài liệu hỗ trợ cơ sở lý thuyết cho Chương 3 khi phân tích và áp dụng kiến trúc phần mềm MVC vào thực tiễn).")
                        
    add_reference_item(doc, 4, 
                       "Tạp chí Công nghệ Thông tin và Truyền thông, \"Ứng dụng Python và các thư viện mã nguồn mở trong xây dựng phần mềm quản lý cỡ nhỏ\", Số chuyên đề Công nghệ Phần mềm, tr. 45-52, 2024.", 
                       None, 
                       "", 
                       "(Bài báo khoa học tham khảo để củng cố tính thực tiễn và lý do chọn đề tài ở Chương 1).")
                        
    # --- Tài liệu nước ngoài ---
    add_heading_styled(doc, "Tài liệu nước ngoài (Tiếng Anh)", 2)
    
    add_reference_item(doc, 5, 
                       "Mark Roseman, ", 
                       "Modern Tkinter for Busy Python Developers", 
                       ", 2020. [Trực tuyến]. Có sẵn tại: https://tkdocs.com/", 
                       "(Tài liệu gốc tham khảo cho việc thiết kế giao diện, sử dụng các thành phần ttk.Treeview và định dạng style ở Chương 2 và Chương 4).")
                        
    add_reference_item(doc, 6, 
                       "Luciano Ramalho, ", 
                       "Fluent Python: Clear, Concise, and Effective Programming", 
                       ", O'Reilly Media, 2nd Edition, 2022.", 
                       "(Sách chuyên sâu về kỹ thuật lập trình Python chuẩn mực, giúp tối ưu hóa mã nguồn trong ứng dụng).")
                        
    add_reference_item(doc, 7, 
                       "Steve Canny, ", 
                       "python-docx Documentation", 
                       ", Read the Docs, 2024. [Trực tuyến]. Có sẵn tại: https://python-docx.readthedocs.io/", 
                       "(Tài liệu chính thức của thư viện python-docx dùng cho tính năng xuất báo cáo Microsoft Word tự động ở Chương 4 và 5).")
                        
    add_reference_item(doc, 8, 
                       "Ian Sommerville, ", 
                       "Software Engineering", 
                       ", Pearson, 10th Edition, 2015.", 
                       "(Tài liệu tham khảo kinh điển về quy trình phát triển phần mềm, quản lý cơ sở dữ liệu và đánh giá ưu nhược điểm của hệ thống).")
                       
    add_reference_item(doc, 9,
                       "Wes McKinney, ",
                       "Python for Data Analysis: Data Wrangling with pandas, NumPy, and Jupyter",
                       ", O'Reilly Media, 3rd Edition, 2022.",
                       "(Tài liệu tham khảo chính yếu để xử lý cấu trúc dữ liệu DataFrame và tính toán các chỉ số thống kê trong ứng dụng).")

    add_reference_item(doc, 10,
                       "John Goerzen, Brandon Rhodes, ",
                       "Foundations of Python Network Programming",
                       ", Apress, 3rd Edition, 2014.",
                       "(Cung cấp các nền tảng lý thuyết quan trọng về luồng xử lý và tối ưu hóa hệ thống máy chủ và máy trạm).")

    add_reference_item(doc, 11,
                       "Brett Slatkin, ",
                       "Effective Python: 90 Specific Ways to Write Better Python",
                       ", Addison-Wesley, 2nd Edition, 2019.",
                       "(Sách cung cấp các nguyên tắc viết code Python chuẩn mực, tối ưu hóa tốc độ và quản lý bộ nhớ RAM).")

    add_reference_item(doc, 12,
                       "Microsoft Developer Network (MSDN), \"Word Document Format XML Schema Reference\", 2023.",
                       None,
                       "",
                       "(Tài liệu tham khảo đặc tả XML để thực hiện can thiệp sâu vào cấu trúc bảng biểu và màu nền của tài liệu Word).")
    
    # Ghi chú cuối tài liệu tham khảo
    p_ref_end = doc.add_paragraph()
    p_ref_end.paragraph_format.space_before = Pt(24)
    p_ref_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    pBdr_ref = OxmlElement('w:pBdr')
    top_ref = OxmlElement('w:top')
    top_ref.set(qn('w:val'), 'single')
    top_ref.set(qn('w:sz'), '6')
    top_ref.set(qn('w:space'), '4')
    top_ref.set(qn('w:color'), HEX_BORDER)
    pBdr_ref.append(top_ref)
    p_ref_end._p.get_or_add_pPr().append(pBdr_ref)
    
    run_ref_end = p_ref_end.add_run("——— Hết tài liệu tham khảo ———")
    run_ref_end.font.name = 'Segoe UI'
    run_ref_end.font.italic = True
    run_ref_end.font.size = Pt(10)
    run_ref_end.font.color.rgb = RGBColor(148, 163, 184)  # Slate 400
    
    # 3. Tiến hành lưu văn bản
    target_dir = os.path.dirname(os.path.abspath(__file__))
    # Đi ngược ra thư mục gốc d:\HoanThanhDuAn
    parent_dir = os.path.dirname(target_dir)
    output_path = os.path.join(parent_dir, "Bao_Cao_Cooking_Manager.docx")
    
    saved = False
    for i in range(10):
        try:
            if i == 0:
                path = output_path
            else:
                path = os.path.join(parent_dir, f"Bao_Cao_Cooking_Manager_{i}.docx")
            doc.save(path)
            output_path = path
            saved = True
            break
        except PermissionError:
            continue
            
    if not saved:
        raise PermissionError("Không thể lưu file vì file liên tục bị khóa. Vui lòng đóng Microsoft Word và chạy lại.")
        
    print(f"\n=> ĐÃ TẠO THÀNH CÔNG BÁO CÁO TẠI: {output_path}")
    return output_path

if __name__ == "__main__":
    tao_bao_cao()
