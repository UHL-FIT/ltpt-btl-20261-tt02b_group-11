# -*- coding: utf-8 -*-
"""
File: them_trang_va_chu_thich.py
Mục đích: Đọc tệp Word hiện có (đã được người dùng sửa), tự động thêm số trang ở giữa in đậm
và thêm chú thích tương ứng dưới 7 hình ảnh đã chèn.
"""

import os
import sys

try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn, nsdecls
except ImportError:
    print("Đang cài đặt python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn, nsdecls

def setup_page_numbering(doc):
    """Thiết lập Header và Footer in đậm số trang ở giữa, ẩn ở trang bìa."""
    for section in doc.sections:
        # Ẩn ở trang bìa
        section.different_first_page_header_footer = True
        
        # --- HEADER (Tiêu đề đầu trang) ---
        header = section.header
        header.is_linked_to_previous = False
        # Xóa header cũ
        for para in header.paragraphs:
            p = para._p
            p.getparent().remove(p)
            
        p_hdr = header.add_paragraph()
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_hdr.paragraph_format.space_after = Pt(4)
        run_hdr = p_hdr.add_run("BÁO CÁO DỰ ÁN: COOKING MANAGER APP")
        run_hdr.font.name = 'Segoe UI'
        run_hdr.font.size = Pt(9)
        run_hdr.font.italic = True
        run_hdr.font.color.rgb = RGBColor(100, 116, 139) # Gray/Slate
        
        # Thêm đường kẻ dưới header
        pBdr = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '6')
        bottom_border.set(qn('w:space'), '4')
        bottom_border.set(qn('w:color'), 'CBD5E1')
        pBdr.append(bottom_border)
        p_hdr._p.get_or_add_pPr().append(pBdr)
        
        # --- FOOTER (Số trang in đậm ở giữa) ---
        footer = section.footer
        footer.is_linked_to_previous = False
        # Xóa footer cũ
        for para in footer.paragraphs:
            p = para._p
            p.getparent().remove(p)
            
        p_ftr = footer.add_paragraph()
        p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # "Trang " in đậm màu tối
        run_ftr_prefix = p_ftr.add_run("Trang ")
        run_ftr_prefix.font.name = 'Segoe UI'
        run_ftr_prefix.font.size = Pt(9.5)
        run_ftr_prefix.font.bold = True
        run_ftr_prefix.font.color.rgb = RGBColor(15, 23, 42) # Dark Slate
        
        # PAGE field
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        
        run_page = p_ftr.add_run()
        run_page.font.name = 'Segoe UI'
        run_page.font.size = Pt(9.5)
        run_page.font.bold = True
        run_page.font.color.rgb = RGBColor(15, 23, 42)
        run_page._r.append(fldChar_begin)
        run_page._r.append(instrText)
        run_page._r.append(fldChar_end)
        
        # Tên phần mềm in nhạt bên cạnh
        run_ftr_suffix = p_ftr.add_run(" | Quản Lý Công Thức Nấu Ăn - Cooking Manager")
        run_ftr_suffix.font.name = 'Segoe UI'
        run_ftr_suffix.font.size = Pt(9)
        run_ftr_suffix.font.color.rgb = RGBColor(100, 116, 139)
        
        # Đường kẻ trên footer
        pBdr_ftr = OxmlElement('w:pBdr')
        top_border = OxmlElement('w:top')
        top_border.set(qn('w:val'), 'single')
        top_border.set(qn('w:sz'), '6')
        top_border.set(qn('w:space'), '4')
        top_border.set(qn('w:color'), 'CBD5E1')
        pBdr_ftr.append(top_border)
        p_ftr._p.get_or_add_pPr().append(pBdr_ftr)

def insert_paragraph_after(paragraph, text="", style=None):
    """Chèn một đoạn văn mới ngay sau đoạn văn hiện tại."""
    new_p_element = OxmlElement('w:p')
    paragraph._p.addnext(new_p_element)
    new_paragraph = docx.text.paragraph.Paragraph(new_p_element, paragraph._parent)
    if text:
        new_paragraph.text = text
    if style:
        new_paragraph.style = style
    return new_paragraph

def add_image_captions(doc):
    """Tìm các hình ảnh trong tài liệu và chèn chú thích tương ứng ngay dưới chúng."""
    captions = [
        "Hình 1. Logo Trường Đại Học Hạ Long",
        "Hình 2. Giao diện chính của ứng dụng Cooking Manager với Sidebar và Danh sách công thức",
        "Hình 3. Giao diện chức năng Nhập liệu và Thêm mới/Lưu cập nhật công thức",
        "Hình 4. Giao diện chức năng tìm kiếm và bộ lọc kết hợp đa tiêu chí",
        "Hình 5. Giao diện màn hình Thống kê tổng quan nhanh trong ứng dụng",
        "Hình 6. Giao diện popup Bảng Chi Tiết hiển thị phân tích thống kê nâng cao với Pandas",
        "Hình 7. Giao diện thông báo xuất file báo cáo Word (.docx) thành công"
    ]
    
    image_count = 0
    paragraphs_to_process = list(doc.paragraphs)
    
    for idx, p in enumerate(paragraphs_to_process):
        # Kiểm tra xem đoạn văn có chứa hình ảnh không (w:drawing hoặc w:graphic)
        if 'w:drawing' in p._element.xml or 'w:graphic' in p._element.xml:
            if image_count < len(captions):
                caption_text = captions[image_count]
                image_count += 1
                
                # Tạo đoạn văn chú thích mới ngay dưới ảnh
                p_cap = insert_paragraph_after(p, caption_text)
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_before = Pt(4)
                p_cap.paragraph_format.space_after = Pt(14)
                
                # Định dạng chữ chú thích: Segoe UI, 9.5pt, Nghiêng, màu xám Slate
                for run in p_cap.runs:
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(9.5)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
                    
                print(f"Đã thêm chú thích: {caption_text}")

def main():
    target_dir = os.path.dirname(os.path.abspath(__file__))
    parent_dir = os.path.dirname(target_dir)
    file_path = os.path.join(parent_dir, "Bao_Cao_Cooking_Manager.docx")
    
    if not os.path.exists(file_path):
        print(f"Lỗi: Không tìm thấy file báo cáo tại: {file_path}")
        return
        
    print(f"Đang xử lý tệp tin: {file_path}...")
    doc = Document(file_path)
    
    # 1. Thêm số trang và header/footer
    setup_page_numbering(doc)
    print("Đã cài đặt định dạng đánh số trang in đậm ở giữa.")
    
    # 2. Thêm các chú thích hình ảnh
    add_image_captions(doc)
    
    # 3. Lưu tài liệu (Hỗ trợ tự động fallback nếu bị khóa)
    saved = False
    output_path = file_path
    for i in range(10):
        try:
            if i == 0:
                path = output_path
            else:
                path = os.path.join(parent_dir, f"Bao_Cao_Cooking_Manager_{i}.docx")
            doc.save(path)
            output_path = path
            saved = True
            break
        except PermissionError:
            continue
            
    if not saved:
        print("Lỗi: Không thể lưu file vì tệp tin đang bị Microsoft Word mở khóa.")
        sys.exit(1)
        
    print(f"\n=> XỬ LÝ THÀNH CÔNG! Báo cáo mới đã được lưu tại: {output_path}")

if __name__ == "__main__":
    main()
